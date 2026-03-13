import io
import json
import os
import re
import tempfile
from datetime import datetime

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from src.externals.diagram.generator import diagram_generator
from src.utils.logger import logger


DEFAULT_THEME_COLOR = "1F4E79"
BODY_TEXT_COLOR = "1F2937"
MUTED_TEXT_COLOR = "6B7280"
SOFT_SURFACE_COLOR = "F8FAFC"
PAGE_BORDER_COLOR = "333333"


class DOCXGenerator:
    def __init__(self, theme_color: str | None = None):
        self.doc = Document()
        self.figure_count = 1
        self.temp_dir = tempfile.TemporaryDirectory()
        self.is_first_section = True
        self.document_title = "Technical Documentation"
        self.theme_color = self._normalize_hex_color(theme_color)
        self.theme_dark = self._mix_hex_color(self.theme_color, "111827", 0.25)
        self.theme_light = self._mix_hex_color(self.theme_color, "FFFFFF", 0.84)
        self._configure_document_styles()

    @staticmethod
    def _normalize_hex_color(value: str | None) -> str:
        cleaned = (value or DEFAULT_THEME_COLOR).strip().lstrip("#").upper()
        if re.fullmatch(r"[0-9A-F]{6}", cleaned):
            return cleaned
        return DEFAULT_THEME_COLOR

    @staticmethod
    def _mix_hex_color(source: str, target: str, ratio: float) -> str:
        source_rgb = [int(source[i:i + 2], 16) for i in range(0, 6, 2)]
        target_rgb = [int(target[i:i + 2], 16) for i in range(0, 6, 2)]
        blended = [
            round(src + (dst - src) * ratio)
            for src, dst in zip(source_rgb, target_rgb)
        ]
        return "".join(f"{value:02X}" for value in blended)

    @staticmethod
    def _rgb(hex_color: str) -> RGBColor:
        return RGBColor.from_string(hex_color)

    def _get_or_add_style(self, name: str, base_style: str = "Normal"):
        styles = self.doc.styles
        try:
            style = styles[name]
        except KeyError:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = styles[base_style]
        return style

    def _set_style_font(
        self,
        style,
        *,
        font_name: str,
        size: float,
        color: str,
        bold: bool = False,
        italic: bool = False,
    ) -> None:
        style.font.name = font_name
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.italic = italic
        style.font.color.rgb = self._rgb(color)

        r_pr = style.element.get_or_add_rPr()
        r_fonts = r_pr.find(qn("w:rFonts"))
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.insert(0, r_fonts)
        for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
            r_fonts.set(qn(f"w:{attr}"), font_name)

    def _configure_document_styles(self) -> None:
        styles = self.doc.styles

        normal = styles["Normal"]
        self._set_style_font(normal, font_name="Calibri", size=11, color=BODY_TEXT_COLOR)
        normal.paragraph_format.space_after = Pt(8)
        normal.paragraph_format.line_spacing = 1.15

        title = styles["Title"]
        self._set_style_font(title, font_name="Cambria", size=28, color=self.theme_color, bold=True)
        title.paragraph_format.space_after = Pt(10)

        for level, size in ((1, 18), (2, 14), (3, 12.5)):
            style = styles[f"Heading {level}"]
            self._set_style_font(
                style,
                font_name="Cambria",
                size=size,
                color=self.theme_color if level == 1 else self.theme_dark,
                bold=True,
            )
            style.paragraph_format.space_before = Pt(18 if level == 1 else 12)
            style.paragraph_format.space_after = Pt(6)
            style.paragraph_format.keep_with_next = True

        for style_name in ("List Bullet", "List Number"):
            style = styles[style_name]
            self._set_style_font(style, font_name="Calibri", size=11, color=BODY_TEXT_COLOR)
            style.paragraph_format.space_after = Pt(4)

        self.subtitle_style_name = "Forge Subtitle"
        subtitle = self._get_or_add_style(self.subtitle_style_name)
        self._set_style_font(subtitle, font_name="Calibri", size=12, color=MUTED_TEXT_COLOR)
        subtitle.paragraph_format.space_after = Pt(6)

        self.summary_style_name = "Forge Summary"
        summary = self._get_or_add_style(self.summary_style_name)
        self._set_style_font(summary, font_name="Calibri", size=11, color=BODY_TEXT_COLOR)
        summary.paragraph_format.space_after = Pt(10)

        self.detail_style_name = "Forge Detail"
        detail = self._get_or_add_style(self.detail_style_name)
        self._set_style_font(detail, font_name="Calibri", size=9.5, color=MUTED_TEXT_COLOR)
        detail.paragraph_format.space_after = Pt(4)

        self.caption_style_name = "Forge Caption"
        caption = self._get_or_add_style(self.caption_style_name)
        self._set_style_font(caption, font_name="Calibri", size=9, color=MUTED_TEXT_COLOR, italic=True)
        caption.paragraph_format.space_before = Pt(4)
        caption.paragraph_format.space_after = Pt(10)

        self.code_style_name = "Forge Code"
        code = self._get_or_add_style(self.code_style_name)
        self._set_style_font(code, font_name="Consolas", size=9.5, color=BODY_TEXT_COLOR)
        code.paragraph_format.left_indent = Inches(0.2)
        code.paragraph_format.right_indent = Inches(0.2)
        code.paragraph_format.space_before = Pt(6)
        code.paragraph_format.space_after = Pt(8)
        code.paragraph_format.line_spacing = 1.0

        self.footer_style_name = "Forge Footer"
        footer = self._get_or_add_style(self.footer_style_name)
        self._set_style_font(footer, font_name="Calibri", size=9, color=MUTED_TEXT_COLOR)
        footer.paragraph_format.space_after = Pt(0)

    def _configure_section_layout(self, section) -> None:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)
        section.header_distance = Inches(0.35)
        section.footer_distance = Inches(0.35)

    def _set_page_borders(self, section) -> None:
        sectPr = section._sectPr
        pgBorders = sectPr.find(qn("w:pgBorders"))
        if pgBorders is not None:
            sectPr.remove(pgBorders)

        pgBorders = OxmlElement("w:pgBorders")
        pgBorders.set(qn("w:offsetFrom"), "page")
        pgBorders.set(qn("w:display"), "allPages")
        pgBorders.set(qn("w:zOrder"), "front")

        for border_name in ["top", "left", "bottom", "right"]:
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "24")
            border.set(qn("w:space"), "24")
            border.set(qn("w:color"), PAGE_BORDER_COLOR)
            pgBorders.append(border)

        following_tags = {
            qn("w:lnNumType"),
            qn("w:pgNumType"),
            qn("w:cols"),
            qn("w:formProt"),
            qn("w:vAlign"),
            qn("w:noEndnote"),
            qn("w:titlePg"),
            qn("w:textDirection"),
            qn("w:bidi"),
            qn("w:rtlGutter"),
            qn("w:docGrid"),
            qn("w:printerSettings"),
            qn("w:sectPrChange"),
        }
        insert_before = next((child for child in sectPr if child.tag in following_tags), None)
        if insert_before is not None:
            insert_before.addprevious(pgBorders)
        else:
            sectPr.append(pgBorders)

    def _set_page_numbers(self, section) -> None:
        footer = section.footer
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.style = self.doc.styles[self.footer_style_name]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self._clear_paragraph_runs(para)

        prefix = para.add_run("Page ")
        prefix.font.color.rgb = self._rgb(MUTED_TEXT_COLOR)
        prefix.font.size = Pt(9)

        run = para.add_run()
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")

        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = "PAGE"

        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "separate")

        t = OxmlElement("w:t")
        t.text = "1"

        fldChar3 = OxmlElement("w:fldChar")
        fldChar3.set(qn("w:fldCharType"), "end")

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(t)
        run._r.append(fldChar3)
        run.font.color.rgb = self._rgb(self.theme_dark)
        run.font.size = Pt(9)

    def _clear_paragraph_runs(self, paragraph) -> None:
        for run in paragraph.runs:
            run_element = run._r
            try:
                run_element.getparent().remove(run_element)
            except Exception:
                pass

    def _set_header(self, section) -> None:
        header = section.header
        paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        paragraph.style = self.doc.styles[self.footer_style_name]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        self._clear_paragraph_runs(paragraph)

        run = paragraph.add_run(f"{self.document_title} | Technical Documentation")
        run.font.color.rgb = self._rgb(self.theme_dark)
        run.font.size = Pt(9)
        self._set_paragraph_bottom_border(paragraph, self.theme_light, size="8")

    def _set_paragraph_bottom_border(self, paragraph, color: str, size: str = "12") -> None:
        p_pr = paragraph._p.get_or_add_pPr()
        p_bdr = p_pr.find(qn("w:pBdr"))
        if p_bdr is None:
            p_bdr = OxmlElement("w:pBdr")
            p_pr.append(p_bdr)

        bottom = p_bdr.find(qn("w:bottom"))
        if bottom is None:
            bottom = OxmlElement("w:bottom")
            p_bdr.append(bottom)

        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), size)
        bottom.set(qn("w:space"), "2")
        bottom.set(qn("w:color"), color)

    def _shade_cell(self, cell, fill: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = tc_pr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            tc_pr.append(shd)
        shd.set(qn("w:fill"), fill)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    def _set_cell_text(self, cell, text: str, *, color: str, bold: bool = False) -> None:
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.style = self.doc.styles["Normal"]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(text)
        run.bold = bold
        run.font.color.rgb = self._rgb(color)

    def _style_table(self, table) -> None:
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        header_fill = self.theme_color

        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(0)
            if row_idx == 0:
                for cell in row.cells:
                    self._shade_cell(cell, header_fill)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.color.rgb = self._rgb("FFFFFF")
            elif row_idx % 2 == 1:
                for cell in row.cells:
                    self._shade_cell(cell, self.theme_light)

    def _add_caption(self, text: str) -> None:
        paragraph = self.doc.add_paragraph(style=self.caption_style_name)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run(text)

    def _trim_text(self, text: str, limit: int = 300) -> str:
        cleaned = re.sub(r"\s+", " ", text or "").strip()
        if len(cleaned) <= limit:
            return cleaned
        return cleaned[: limit - 3].rstrip() + "..."

    def _add_overview_table(self, title: str, description: str, page_count: int | None) -> None:
        table = self.doc.add_table(rows=4, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        generated_on = datetime.now().strftime("%d %b %Y")
        rows = [
            ("Project", title),
            ("Target Length", f"~{page_count} pages" if page_count else "Flexible"),
            ("Generated On", generated_on),
            ("Coverage", self._trim_text(description or "Documentation assembled from the provided project brief.", 180)),
        ]

        for row_idx, (label, value) in enumerate(rows):
            left = table.cell(row_idx, 0)
            right = table.cell(row_idx, 1)
            self._shade_cell(left, self.theme_color)
            self._shade_cell(right, self.theme_light if row_idx % 2 == 0 else SOFT_SURFACE_COLOR)
            self._set_cell_text(left, label, color="FFFFFF", bold=True)
            self._set_cell_text(right, value, color=BODY_TEXT_COLOR)

    def _add_document_outline(self, sections: list[dict[str, str | None]] | None) -> None:
        self.doc.add_heading("Document Structure", level=1)
        intro = self.doc.add_paragraph(style=self.summary_style_name)
        intro.add_run("This document is organized into the following major sections.")

        if not sections:
            empty = self.doc.add_paragraph(style=self.detail_style_name)
            empty.add_run("Detailed sections will be generated dynamically from the project brief.")
            return

        for section in sections:
            outline = self.doc.add_paragraph(style="List Number")
            title_run = outline.add_run(section.get("title") or "Untitled Section")
            title_run.bold = True
            title_run.font.color.rgb = self._rgb(self.theme_dark)

            description = (section.get("description") or "").strip()
            if description:
                detail = self.doc.add_paragraph(style=self.detail_style_name)
                detail.paragraph_format.left_indent = Inches(0.35)
                detail.add_run(self._trim_text(description, 180))

    def _parse_inline_formatting(self, paragraph, text: str) -> None:
        parts = text.split("**")
        for idx, part in enumerate(parts):
            if not part:
                continue
            run = paragraph.add_run(part)
            if idx % 2 == 1:
                run.bold = True

    def _append_word_field(self, paragraph, instruction: str, placeholder_text: str = "") -> None:
        run = paragraph.add_run()

        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")

        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = instruction

        fld_separate = OxmlElement("w:fldChar")
        fld_separate.set(qn("w:fldCharType"), "separate")

        placeholder = OxmlElement("w:t")
        placeholder.text = placeholder_text

        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")

        run._r.append(fld_begin)
        run._r.append(instr_text)
        run._r.append(fld_separate)
        run._r.append(placeholder)
        run._r.append(fld_end)

    def _add_table_of_contents(self) -> None:
        self.doc.add_heading("Contents", level=1)
        paragraph = self.doc.add_paragraph(style=self.summary_style_name)
        self._append_word_field(paragraph, 'TOC \\o "1-3" \\h \\z \\u')
        self.doc.add_paragraph()

    def add_title_page(
        self,
        title: str,
        description: str = "",
        page_count: int | None = None,
        sections: list[dict[str, str | None]] | None = None,
    ) -> None:
        self.document_title = title.strip() or "Technical Documentation"
        self.doc.core_properties.title = title
        self.doc.core_properties.subject = "Technical Project Documentation"
        self.doc.core_properties.author = "LetsCreateDoc"
        self.doc.core_properties.last_modified_by = "LetsCreateDoc"
        self.doc.core_properties.comments = "Generated by LetsCreateDoc"

        for _ in range(4):
            self.doc.add_paragraph()

        title_paragraph = self.doc.add_paragraph(style="Title")
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_paragraph.add_run(title)

        subtitle = self.doc.add_paragraph(style=self.subtitle_style_name)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.add_run("Technical Design and Implementation Document")

        generated = self.doc.add_paragraph(style=self.detail_style_name)
        generated.alignment = WD_ALIGN_PARAGRAPH.CENTER
        generated.add_run(datetime.now().strftime("Generated on %d %b %Y"))

        if description.strip():
            summary = self.doc.add_paragraph(style=self.summary_style_name)
            summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
            summary.add_run(self._trim_text(description, 260))

        accent_rule = self.doc.add_paragraph()
        accent_rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
        accent_rule.add_run(" ")
        self._set_paragraph_bottom_border(accent_rule, self.theme_color, size="18")

        self.doc.add_page_break()

        self.doc.add_heading("Document Overview", level=1)
        overview = self.doc.add_paragraph(style=self.summary_style_name)
        overview.add_run(
            "The front matter below summarizes the project scope, expected depth, and the section map used for the final document."
        )
        self._add_overview_table(title, description, page_count)
        self.doc.add_paragraph()
        self._add_table_of_contents()
        self._add_document_outline(sections)
        self.is_first_section = False

    async def add_markdown_section(self, section_title: str, content: str) -> None:
        if not self.is_first_section:
            self.doc.add_section(WD_SECTION_START.NEW_PAGE)

        if not content.strip().startswith("# "):
            self.doc.add_heading(section_title, level=1)

        def extract_diagram(match):
            json_str = match.group(1)
            try:
                data = json.loads(json_str, strict=False)
                code = data.get("code", "") or data.get("content", "") or data.get("mermaid", "")
                caption = data.get("caption", "")
                if code:
                    return f"\n```mermaid\n{code}\n```\n" + (f"\nCAPTION: {caption}\n" if caption else "")
            except Exception:
                code_match = re.search(r'"(?:code|content|mermaid)"\s*:\s*"(.*?)"\s*\}', json_str, re.DOTALL)
                if code_match:
                    code = code_match.group(1).replace("\\n", "\n").replace('\\"', '"')
                    return f"\n```mermaid\n{code}\n```\n"
            return match.group(0)

        content = re.sub(r"generate_diagram\(\s*(\{.*?\})\s*\)", extract_diagram, content, flags=re.DOTALL)

        lines = content.split("\n")
        i = 0
        in_code_block = False
        code_content = ""
        current_table = []

        def commit_table():
            nonlocal current_table
            if not current_table:
                return

            valid_rows = [row for row in current_table if not re.match(r"^[\|\-\s\:]+$", row)]
            if valid_rows:
                cols = len([item for item in valid_rows[0].split("|") if item.strip()])
                if cols > 0:
                    try:
                        table = self.doc.add_table(rows=len(valid_rows), cols=cols)
                        for row_idx, row_text in enumerate(valid_rows):
                            cells = [item.strip() for item in row_text.split("|") if item.strip()]
                            for col_idx, cell_text in enumerate(cells[:cols]):
                                paragraph = table.cell(row_idx, col_idx).paragraphs[0]
                                self._parse_inline_formatting(paragraph, cell_text)
                        self._style_table(table)
                        self.doc.add_paragraph()
                    except Exception as exc:
                        logger.error(f"Table parsing error: {exc}")
            current_table = []

        while i < len(lines):
            line = lines[i].rstrip()
            stripped = line.strip()

            if in_code_block:
                if stripped.startswith("```"):
                    in_code_block = False
                    if code_content.startswith("mermaid\n"):
                        mermaid_code = code_content[8:]
                        img_bytes = await diagram_generator.generate_image(mermaid_code)
                        if img_bytes:
                            img_path = os.path.join(self.temp_dir.name, f"fig_{self.figure_count}.png")
                            with open(img_path, "wb") as file_handle:
                                file_handle.write(img_bytes)

                            custom_caption = ""
                            if i + 1 < len(lines) and lines[i + 1].startswith("CAPTION:"):
                                custom_caption = lines[i + 1].replace("CAPTION:", "").strip()
                                i += 1

                            caption_text = custom_caption if custom_caption else f"Diagram {self.figure_count}"
                            try:
                                self.doc.add_picture(img_path, width=Inches(6.1))
                            except Exception as exc:
                                logger.error(f"Failed to add picture: {exc}")
                            self._add_caption(caption_text)
                            self.figure_count += 1
                    else:
                        paragraph = self.doc.add_paragraph(style=self.code_style_name)
                        paragraph.add_run(code_content.rstrip())
                        self._set_paragraph_bottom_border(paragraph, self.theme_light, size="6")
                else:
                    code_content += lines[i] + "\n"
                i += 1
                continue

            if stripped.startswith("```"):
                commit_table()
                in_code_block = True
                code_content = stripped[3:] + "\n" if len(stripped) > 3 else ""
                i += 1
                continue

            if stripped.startswith("|"):
                current_table.append(stripped)
                i += 1
                continue

            commit_table()

            if stripped.startswith("#"):
                level = len(stripped.split(" ")[0])
                text = stripped[level:].strip()
                if 1 <= level <= 9:
                    self.doc.add_heading(text, level=min(level, 3))
            elif stripped.startswith("- ") or stripped.startswith("* "):
                paragraph = self.doc.add_paragraph(style="List Bullet")
                self._parse_inline_formatting(paragraph, stripped[2:])
            elif re.match(r"^\d+\.\s", stripped):
                paragraph = self.doc.add_paragraph(style="List Number")
                idx = stripped.find(" ")
                self._parse_inline_formatting(paragraph, stripped[idx + 1 :])
            elif stripped == "":
                pass
            else:
                if stripped.startswith("![") and "](" in stripped:
                    pass
                elif stripped.startswith("CAPTION:"):
                    pass
                else:
                    paragraph = self.doc.add_paragraph()
                    self._parse_inline_formatting(paragraph, stripped)
            i += 1

        commit_table()

    def add_image(self, image_bytes: bytes, caption: str = "") -> None:
        try:
            img_path = os.path.join(self.temp_dir.name, f"ext_fig_{self.figure_count}.png")
            with open(img_path, "wb") as file_handle:
                file_handle.write(image_bytes)

            self.doc.add_picture(img_path, width=Inches(6.1))
            caption_text = caption if caption else f"Diagram {self.figure_count}"
            self._add_caption(caption_text)
            self.figure_count += 1
        except Exception as exc:
            logger.error(f"Failed to add external image: {exc}")

    def get_docx_bytes(self) -> io.BytesIO:
        try:
            for section in self.doc.sections:
                self._configure_section_layout(section)
                self._set_page_borders(section)
                self._set_header(section)
                self._set_page_numbers(section)

            out_path = os.path.join(self.temp_dir.name, "output.docx")
            self.doc.save(out_path)

            with open(out_path, "rb") as file_handle:
                docx_bytes = file_handle.read()

            return io.BytesIO(docx_bytes)
        finally:
            self.temp_dir.cleanup()
